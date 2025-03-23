import streamlit as st
import base64
import json
import requests
import os
from pydub import AudioSegment
from io import BytesIO
import shutil
from docx import Document

def crea_docx_da_risposta(risposta_json):
    doc = Document()
    doc.add_heading("Trascrizione vocale", level=1)

    # Estrai la trascrizione
    for result in risposta_json.get("results", []):
        for alt in result.get("alternatives", []):
            transcript = alt.get("transcript", "")
            doc.add_paragraph(transcript)

    # Salva il file temporaneo
    doc_path = "/tmp/trascrizione.docx"
    doc.save(doc_path)

    # Riapri e restituisci i byte
    with open(doc_path, "rb") as f:
        return f.read()


if not shutil.which("ffmpeg"):
    st.error("❌ FFmpeg non è installato nell'ambiente. Assicurati che `packages.txt` contenga `ffmpeg`.")

st.title("Google Speech-to-Text - FLAC/MP3 Uploader")

# Inserisci la tua API Key qui oppure usa una variabile d'ambiente
api_key_input = st.text_input("🔑 Inserisci la tua Google API Key", type="password")

# Salva nell session_state se presente
if api_key_input:
    st.session_state["api_key"] = api_key_input

# Upload file
uploaded_file = st.file_uploader("Carica un file audio (.flac o .mp3)", type=["flac", "mp3"])

if uploaded_file is not None:
    # Converti mp3 in flac in memoria
    file_extension = uploaded_file.name.split(".")[-1].lower()
    
    if file_extension == "mp3":
        st.info("Convertendo MP3 in FLAC...")
        audio = AudioSegment.from_file(uploaded_file, format="mp3")
    elif file_extension == "flac":
        audio = AudioSegment.from_file(uploaded_file, format="flac")
    else:
        st.error("Formato non supportato.")
        st.stop()

    # Imposta sample rate a 44100 e esporta in memoria come FLAC
    audio = audio.set_frame_rate(44100).set_channels(1)
    flac_buffer = BytesIO()
    audio.export(flac_buffer, format="flac")
    flac_bytes = flac_buffer.getvalue()

    # Codifica in base64
    audio_base64 = base64.b64encode(flac_bytes).decode("utf-8")

    st.success(f"{uploaded_file.name} preparato per l'invio!")

    config = {
        "encoding": "FLAC",
        "sampleRateHertz": 44100,
        "languageCode": "it-IT"  # cambia in en-US se serve
    }

    request_payload = {
        "config": config,
        "audio": {
            "content": audio_base64
        }
    }

    if st.button("Invia a Google Speech-to-Text"):
        api_key = st.session_state.get("api_key", "")
        if not api_key:
            st.error("⚠️ Nessuna API Key trovata. Inseriscila nel campo sopra.")
            st.stop()
        try:
            url = f"https://speech.googleapis.com/v1/speech:recognize?key={api_key}"
            response = requests.post(url, json=request_payload)
            response.raise_for_status()
            result = response.json()

            with open("result.json", "w") as f:
                json.dump(result, f, indent=2)

            # Mostra la risposta
            st.success("Risposta ricevuta!")
            st.json(result)
            
            # Pulsante per scaricare la trascrizione in formato Word
            docx_bytes = crea_docx_da_risposta(result)
            st.download_button(
                label="💾 Scarica trascrizione in formato Word",
                data=docx_bytes,
                file_name="trascrizione.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except requests.exceptions.HTTPError as http_err:
            st.error(f"Errore nella richiesta: {response.status_code} - {response.reason}")
            try:
                st.code(response.json(), language="json")
            except:
                st.text(response.text)

        

